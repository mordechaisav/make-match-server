import io
from pathlib import Path
from typing import Callable

import docx
import pymupdf
import pytesseract
from fastapi import HTTPException, UploadFile, status
from PIL import Image
from pypdf import PdfReader

# Hebrew resumes are often plain photos/scans with no embedded text layer.
# OCR is done locally with Tesseract (deterministic, no rate limits, no
# hallucination risk) rather than an LLM vision model - the recognized text
# is then fed through the existing Groq-based field-extraction pipeline like
# any other text. Deployment must have the `tesseract-ocr` binary plus the
# Hebrew language data (`tesseract-ocr-heb` on Debian/Ubuntu) installed.
_OCR_LANGUAGES = "heb+eng"
_OCR_RENDER_DPI = 300


def _ocr_pdf(data: bytes) -> str:
    parts = []
    with pymupdf.open(stream=data, filetype="pdf") as document:
        for page in document:
            png_bytes = page.get_pixmap(dpi=_OCR_RENDER_DPI).tobytes("png")
            image = Image.open(io.BytesIO(png_bytes))
            parts.append(pytesseract.image_to_string(image, lang=_OCR_LANGUAGES))
    return "\n".join(parts)


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if text.strip():
        return text
    # no embedded text layer - it's a scanned/photographed resume, OCR it
    return _ocr_pdf(data)


def _extract_docx_text(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


_EXTRACTORS_BY_EXTENSION: dict[str, Callable[[bytes], str]] = {
    ".pdf": _extract_pdf_text,
    ".docx": _extract_docx_text,
}


def extract_text_from_upload(file: UploadFile) -> str:
    """Read an uploaded resume (.pdf or .docx) and return its raw text.

    Dispatches purely on the filename extension - content_type is set by the
    client and isn't reliable enough to trust for routing to a parser.
    """
    extension = Path(file.filename or "").suffix.lower()
    extractor = _EXTRACTORS_BY_EXTENSION.get(extension)
    if extractor is None:
        raise HTTPException(
            status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail="Unsupported file type - only .pdf and .docx resumes are accepted",
        )

    data = file.file.read()
    if not data:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty")

    try:
        text = extractor(data)
    except Exception as exc:
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            detail="Could not read the uploaded file - it may be corrupted or password protected",
        ) from exc

    if not text.strip():
        raise HTTPException(status.HTTP_400_BAD_REQUEST, detail="Could not extract any text from the uploaded file")
    return text


def get_text_extractor() -> Callable[[UploadFile], str]:
    return extract_text_from_upload
